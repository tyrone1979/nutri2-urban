#!/usr/bin/env python3
"""Rebuild Supplementary Material.docx with ONLY content cited by main.docx.

Cited by main.docx:
  S1  SHAP main effects
  S2  SHAP interactions
  Fig S1 SHAP summary plot
  S3  Leave-one-province-out
  S4  Feature-set ablation
  S5  Category-specific accuracy across missing rates (Rural & Urban ONLY)
  S6  Hyperparameters
  S7  Binary administrative classification
  S8  Urban probability threshold tuning
  S9  Transitional FatER band sensitivity (Plan B)

Excluded: Plan A leftovers (acc≈0.999), Transitional class-accuracy columns,
          unused baseline/ablation dumps.
"""
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "public" / "Supplementary Material.docx"
BACKUP = ROOT / "public" / "Supplementary Material.docx.bak"
FIG_S1 = ROOT / "figures" / "FigS1_shap_summary.tiff"
FIG_S1_PNG = ROOT / "figures" / "shap_summary.png"


def set_cell(cell, text: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].text = str(text)
        for p in cell.paragraphs[1:]:
            p.text = ""
    else:
        cell.text = str(text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.bold = True


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def fill_table(table, headers, rows) -> None:
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell(table.rows[i].cells[j], val)


def main():
    shutil.copy2(DOCX, BACKUP)

    shap = pd.read_csv(ROOT / "results/shap_main_effects.csv")
    interact = pd.read_csv(ROOT / "results/shap_interactions.csv").head(15)
    spatial = pd.read_csv(ROOT / "results/spatial_validation.csv")
    ablation = pd.read_csv(ROOT / "results/feature_ablation.csv")
    class_acc = pd.read_csv(ROOT / "results/missing_rate_class_acc.csv")
    binary = pd.read_csv(ROOT / "results/binary_classification.csv").iloc[0]
    thr = pd.read_csv(ROOT / "results/threshold_tuning.csv")
    band = pd.read_csv(ROOT / "results/threshold_sensitivity.csv")

    feat_names = {
        "fat_energy_ratio": "Fat energy ratio",
        "carbo_energy_ratio": "Carbohydrate energy ratio",
        "protn_energy_ratio": "Protein energy ratio",
        "fat_carbo_ratio": "Fat-to-carbohydrate ratio",
        "Year": "Survey year",
        "Province": "Province",
    }

    doc = Document()
    title = doc.add_paragraph("Supplementary Material")
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(14)

    # ---- S1 ----
    add_caption(doc, "Supplementary Table S1. SHAP feature importance (mean |SHAP|).")
    add_note(
        doc,
        "Values are global mean absolute SHAP values from TreeExplainer on the held-out test set. "
        "Higher values indicate greater average contribution to model predictions.",
    )
    t1 = doc.add_table(rows=1 + len(shap), cols=2)
    fill_table(
        t1,
        ["Feature", "Mean |SHAP|"],
        [
            [feat_names.get(r["feature"], r["feature"]), f"{r['main_effect']:.3f}"]
            for _, r in shap.iterrows()
        ],
    )

    # ---- S2 ----
    add_caption(doc, "Supplementary Table S2. Top SHAP feature interaction strengths.")
    add_note(
        doc,
        "Interaction strength is the mean absolute SHAP interaction value. "
        "Pairs are ranked by interaction magnitude.",
    )
    t2 = doc.add_table(rows=1 + len(interact), cols=4)
    fill_table(
        t2,
        ["Rank", "Feature 1", "Feature 2", "Interaction strength"],
        [
            [
                str(i),
                feat_names.get(r["feature_1"], r["feature_1"]),
                feat_names.get(r["feature_2"], r["feature_2"]),
                f"{r['interaction_strength']:.4f}",
            ]
            for i, (_, r) in enumerate(interact.iterrows(), start=1)
        ],
    )

    # ---- Fig S1 ----
    add_caption(doc, "Supplementary Figure S1. SHAP summary plot (global feature effects).")
    fig_path = FIG_S1_PNG if FIG_S1_PNG.exists() else (FIG_S1 if FIG_S1.exists() else None)
    if fig_path is not None:
        p = doc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(str(fig_path), width=Inches(5.8))
    add_note(
        doc,
        "Figure S1. SHAP summary plot for Urban-category contributions. Each point is one observation; "
        "horizontal position is the SHAP value (impact on the Urban prediction). Colour encodes the raw "
        "feature value from low (blue) to high (red). Points are vertically jittered within each feature "
        "row to reduce overplotting (not ordered by feature value).",
    )

    # ---- S3 ----
    mean_acc = spatial["Accuracy"].mean()
    sd_acc = spatial["Accuracy"].std()
    mean_f1 = spatial["Macro_F1"].mean()
    sd_f1 = spatial["Macro_F1"].std()
    add_caption(
        doc,
        f"Supplementary Table S3. Leave-one-province-out spatial validation "
        f"(mean accuracy {mean_acc:.3f}, SD {sd_acc:.3f}).",
    )
    t3 = doc.add_table(rows=2 + len(spatial), cols=4)
    rows3 = [
        [r["Province"], f"{int(r['n_test']):,}", f"{r['Accuracy']:.3f}", f"{r['Macro_F1']:.3f}"]
        for _, r in spatial.iterrows()
    ]
    rows3.append(["Mean ± SD", "", f"{mean_acc:.3f} ± {sd_acc:.3f}", f"{mean_f1:.3f} ± {sd_f1:.3f}"])
    fill_table(t3, ["Province", "N (test)", "Accuracy", "Macro-F1"], rows3)

    # ---- S4 ----
    add_caption(
        doc,
        "Supplementary Table S4. Feature-set ablation (T2 + transitional overlay labels).",
    )
    add_note(
        doc,
        "Primary interpretation should emphasise the binary Rural/Urban task (Table S7). "
        "Excluding FatER yields only a modest drop under the three-class descriptive framework.",
    )
    label_map = {
        "full": "Full (6 features)",
        "no_fater": "No FatER (5 features)",
        "nutrients_only": "Macronutrients only (4 features)",
        "spatiotemporal_only": "Year + Province only (2 features)",
    }
    t4 = doc.add_table(rows=1 + len(ablation), cols=4)
    fill_table(
        t4,
        ["Feature set", "Accuracy", "Macro-F1", "Cohen's κ"],
        [
            [
                label_map.get(r["Feature_Set"], r["Feature_Set"]),
                f"{r['Accuracy']:.3f}",
                f"{r['Macro_F1']:.3f}",
                f"{r['Kappa']:.3f}",
            ]
            for _, r in ablation.iterrows()
        ],
    )

    # ---- S5: Rural & Urban ONLY (no Transitional 0.99) ----
    add_caption(
        doc,
        "Supplementary Table S5. Category-specific accuracy across missing rates "
        "(Rural and Urban only).",
    )
    add_note(
        doc,
        "Transitional accuracy is omitted because that stratum is partially defined by FatER "
        "and yields near-perfect scores that are not informative for administrative label recovery. "
        "Urban accuracy remains the challenging stratum (~0.41–0.43).",
    )
    rates = sorted(class_acc["missing_rate"].unique())
    t5 = doc.add_table(rows=1 + len(rates), cols=3)
    rows5 = []
    for rate in rates:
        sub = class_acc.loc[class_acc["missing_rate"] == rate]
        rural = float(sub.loc[sub["class"] == "Rural", "proposed_acc"].iloc[0])
        urban = float(sub.loc[sub["class"] == "Urban", "proposed_acc"].iloc[0])
        rows5.append([f"{int(rate * 100)}%", f"{rural:.3f}", f"{urban:.3f}"])
    fill_table(t5, ["Missing rate", "Rural accuracy", "Urban accuracy"], rows5)

    # ---- S6 ----
    add_caption(doc, "Supplementary Table S6. Hyperparameter configuration for Balanced XGBoost.")
    add_note(doc, "Selected via grid search with five-fold cross-validation on the training set.")
    t6 = doc.add_table(rows=10, cols=3)
    fill_table(
        t6,
        ["Parameter", "Value", "Purpose"],
        [
            ["Boosting rounds", "600", "Extended fitting with early stopping via CV"],
            ["Maximum tree depth", "4", "Control complexity / overfitting"],
            ["Learning rate", "0.05", "Shrinkage for stable updates"],
            ["Subsample ratio", "0.9", "Row subsampling"],
            ["Column subsample", "0.9", "Feature subsampling"],
            ["Min child weight", "5", "Leaf regularity"],
            ["Gamma", "0.2", "Minimum loss reduction for splits"],
            ["L1 / L2 regularisation", "0.5 / 1.2", "Shrinkage of leaf weights"],
            ["Class weighting", "Balanced", "Address Rural/Urban/Transitional imbalance"],
        ],
    )

    # ---- S7 ----
    add_caption(
        doc,
        "Supplementary Table S7. Binary administrative classification "
        "(Rural vs Urban, excluding Transitional) under 30% simulated missingness.",
    )
    add_note(
        doc,
        "This is the primary inferential target for non-circular evaluation of administrative label recovery.",
    )
    t7 = doc.add_table(rows=2, cols=8)
    fill_table(
        t7,
        [
            "Scenario", "N (test)", "Accuracy", "Macro-F1", "Weighted-F1",
            "Urban recall", "Rural recall", "Rural PPV",
        ],
        [[
            "Binary XGBoost",
            str(int(binary["n_test"])),
            f"{binary['accuracy_masked']:.3f}",
            f"{binary['macro_f1_masked']:.3f}",
            f"{binary['weighted_f1_masked']:.3f}",
            f"{binary['urban_recall_masked']:.3f}",
            f"{binary['rural_recall_masked']:.3f}",
            f"{binary['rural_precision_masked']:.3f}",
        ]],
    )

    # ---- S8 ----
    add_caption(
        doc,
        "Supplementary Table S8. Urban probability threshold tuning "
        "(three-class model, 30% masked labels).",
    )
    t8 = doc.add_table(rows=1 + len(thr), cols=5)
    fill_table(
        t8,
        ["Urban threshold", "Overall accuracy", "Macro-F1", "Urban recall", "Rural recall"],
        [
            [
                f"{r['urban_threshold']:.2f}",
                f"{r['overall_acc']:.3f}",
                f"{r['macro_f1']:.3f}",
                f"{r['urban_recall']:.3f}",
                f"{r['rural_recall']:.3f}",
            ]
            for _, r in thr.iterrows()
        ],
    )

    # ---- S9: Plan B band sensitivity (realistic, not ~1.0) ----
    add_caption(
        doc,
        "Supplementary Table S9. Transitional FatER band sensitivity "
        "(T2 base labels with varying transitional overlay).",
    )
    add_note(
        doc,
        "Accuracy changes with band width because transitional prevalence and label composition change; "
        "these are descriptive sensitivity checks, not claims of near-perfect prediction.",
    )
    t9 = doc.add_table(rows=1 + len(band), cols=6)
    fill_table(
        t9,
        ["Trans low", "Trans high", "Accuracy", "Macro-F1", "Kappa", "Transitional prevalence"],
        [
            [
                f"{r['trans_low']:.2f}",
                f"{r['trans_high']:.2f}",
                f"{r['accuracy']:.3f}",
                f"{r['macro_f1']:.3f}",
                f"{r['kappa']:.3f}",
                f"{r['transitional_prev']:.3f}",
            ]
            for _, r in band.iterrows()
        ],
    )

    doc.save(str(DOCX))
    print(f"Rebuilt {DOCX}")
    print("Kept: S1–S9 + Fig S1 (Rural/Urban only in S5; no Plan A 0.999 tables)")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
