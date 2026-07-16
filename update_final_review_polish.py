#!/usr/bin/env python3
"""Final-review polish: protocol guide, MAR z() note, inflation theory, S9 wording, formatting."""
import re
from pathlib import Path

from docx import Document

MAIN = Path(__file__).resolve().parent / "public" / "main.docx"


def main():
    doc = Document(str(MAIN))

    guide = (
        "Protocol application guide. Researchers adapting this blueprint to another survey should: "
        "(i) choose a reference cohort with complete contextual labels for internal validation; "
        "(ii) specify MCAR masking rates matching expected incompleteness and at least one MAR mechanism "
        "tied to observed covariates with reported propensity coefficients, and report the resulting "
        "realized missing rate alongside the propensity coefficient (e.g., β₁ = 2.0 yielded a realized "
        "rate of 30%); "
        "(iii) pre-declare a primary estimand that is not defined by the same features used as predictors—"
        "external or administrative labels are strongly preferred as the primary target variable to minimise "
        "definitional overlap of the kind encountered when transitional strata are thresholded on predictors; "
        "(iv) evaluate both predictive metrics and at least one downstream target parameter "
        "(e.g., an adjusted exposure coefficient) under true versus imputed labels; "
        "(v) assess temporal and/or spatial holdout generalisability; and "
        "(vi) report calibration if probabilities will be reused as weights or multiple-imputation draws. "
        "Hyperparameters and comparator imputation methods can be swapped without changing the protocol skeleton."
    )

    fragility = (
        "Threshold fragility of the three-class descriptive labels. Varying the transitional FatER band "
        "changed three-class accuracy sharply: under the default band [0.23, 0.30] accuracy was 0.782, "
        "whereas widening to [0.25, 0.32] reduced accuracy to 0.679 (Supplementary Table S9). "
        "This large swing, driven solely by a 2-percentage-point shift in a threshold applied to a predictor "
        "variable, renders the three-class accuracy fundamentally non-interpretable as a measure of model "
        "performance and provides the decisive justification for focusing all primary conclusions on the "
        "binary Rural/Urban task (masked accuracy 0.716; Supplementary Table S7)."
    )

    inflation = (
        "Analyses using imputed labels preserved the direction and statistical significance of "
        "urban–rural differences across all dietary indicators (Table 8; Figure 5). Cohen's d was "
        "moderately inflated (25% relative increase on average) when using imputed labels, because "
        "Transitional observations were assigned to Rural or Urban groups, sharpening group contrasts. "
        "In measurement-error terms, this behaves like non-differential misclassification of the "
        "contextual exposure with respect to the dietary outcomes: reassignment changes group means "
        "and increases between-group separation, thereby inflating effect sizes even when outcome "
        "measurement is unchanged. Statistical significance was preserved in all cases; imputed labels "
        "are better suited to descriptive comparisons or covariate adjustment than standalone effect estimation."
    )

    to_remove = []
    for p in doc.paragraphs:
        t = p.text

        if t.startswith("Protocol application guide"):
            p.text = guide

        elif t.startswith("MCAR. On the held-out test set"):
            newt = re.sub(
                r"where z\(FatER\)[^,]+,",
                "where z(FatER) denotes FatER standardised on the held-out test set "
                "(mean and SD computed on that test set only, matching the masking experiment),",
                t,
                count=1,
            )
            if newt == t and "mean and SD computed on that test set" not in t:
                newt = t.replace(
                    "standardised FatER on the test set",
                    "FatER standardised on the held-out test set "
                    "(mean and SD computed on that test set only, matching the masking experiment)",
                )
            p.text = newt

        elif t.startswith("Threshold fragility of the three-class"):
            p.text = fragility

        elif (
            t.startswith("Analyses using imputed labels preserved the direction")
            and "Cohen's d was moderately inflated" in t
        ):
            p.text = inflation

        elif t.strip() in {"### Strengths", "### Strengths "}:
            p.text = "Strengths"
            for r in p.runs:
                r.bold = True

        elif t.strip() in {"### Limitations", "### Limitations "}:
            p.text = "Limitations"
            for r in p.runs:
                r.bold = True

        elif t.startswith("Second, the framework was developed using data from a single country"):
            to_remove.append(p)

        elif t.startswith("Inferred labels are probabilistic proxies"):
            p.text = t.replace("Inferred labels", "Imputed labels", 1)

    for p in doc.paragraphs:
        if "inferred labels" in p.text:
            p.text = p.text.replace("inferred labels", "imputed labels")
        if "Inferred labels" in p.text:
            p.text = p.text.replace("Inferred labels", "Imputed labels")

    for p in to_remove:
        p._element.getparent().remove(p._element)

    for p in doc.paragraphs:
        if p.text.strip() in {"Strengths", "Limitations"}:
            for r in p.runs:
                r.bold = True

    doc.save(str(MAIN))

    # Verify
    doc = Document(str(MAIN))
    ok = {
        "guide_realized": any("realized rate of 30" in p.text for p in doc.paragraphs),
        "guide_admin": any("external or administrative labels are strongly preferred" in p.text for p in doc.paragraphs),
        "z_test": any("mean and SD computed on that test set only" in p.text for p in doc.paragraphs),
        "nondiff": any("non-differential misclassification" in p.text for p in doc.paragraphs),
        "fragility": any("fundamentally non-interpretable" in p.text for p in doc.paragraphs),
        "no_hash_strengths": not any(p.text.strip() == "### Strengths" for p in doc.paragraphs),
        "no_hash_lim": not any(p.text.strip() == "### Limitations" for p in doc.paragraphs),
        "has_strengths": any(p.text.strip() == "Strengths" for p in doc.paragraphs),
        "no_dup_second": not any(
            p.text.startswith("Second, the framework was developed using data from a single country")
            for p in doc.paragraphs
        ),
        "no_inferred_labels": not any("inferred labels" in p.text for p in doc.paragraphs),
    }
    for k, v in ok.items():
        print(("OK" if v else "FAIL"), k)
    print(f"Updated {MAIN}")


if __name__ == "__main__":
    main()
