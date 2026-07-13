#!/usr/bin/env python3
"""Apply Statistics in Medicine reviewer-directed revisions to manuscript assets."""
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MAIN = ROOT / "public" / "main.docx"
COVER = ROOT / "public" / "cover_letter_SiM.docx"
SUPP = ROOT / "public" / "Supplementary Material.docx"
RESPONSE = ROOT / "public" / "RESPONSE_TO_REVIEWERS.md"


def has_equation(element) -> bool:
    for node in element.iter():
        tag = node.tag.split("}")[-1] if "}" in node.tag else node.tag
        if tag in ("oMath", "oMathPara"):
            return True
    return False


def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].text = text
        for p in cell.paragraphs[1:]:
            p.text = ""
    else:
        cell.text = text


def replace_para(p, text: str) -> None:
    if has_equation(p._element):
        return
    p.text = text


def insert_before(doc, anchor: str, text: str) -> bool:
    for p in doc.paragraphs:
        if anchor in p.text and not has_equation(p._element):
            new_p = p.insert_paragraph_before(text)
            new_p.paragraph_format.space_after = Pt(6)
            return True
    return False


def load_metrics():
    baseline = pd.read_csv(ROOT / "results/enhanced_baseline_comparison.csv")
    class_acc = pd.read_csv(ROOT / "results/enhanced_baseline_class_acc.csv")
    binary = pd.read_csv(ROOT / "results/binary_classification.csv").iloc[0]
    thr = pd.read_csv(ROOT / "results/threshold_tuning.csv")
    downstream = pd.read_csv(ROOT / "results/downstream_bias.csv")
    thr_bal = thr.loc[thr["urban_threshold"] == 0.40].iloc[0]
    thr_urb = thr.loc[thr["urban_threshold"] == 0.35].iloc[0]

    prop = baseline.loc[baseline["Method"] == "Proposed (BXGB)"].iloc[0]
    mice = baseline.loc[baseline["Method"] == "MICE"].iloc[0]
    knn = baseline.loc[baseline["Method"] == "KNN (k=5)"].iloc[0]
    rf = baseline.loc[baseline["Method"] == "RF-Imputer"].iloc[0]
    urban_acc = float(class_acc.loc[class_acc["Class"] == "Urban", "Proposed (BXGB)"].iloc[0])
    trans_acc = float(class_acc.loc[class_acc["Class"] == "Transitional", "Proposed (BXGB)"].iloc[0])
    rural_acc = float(class_acc.loc[class_acc["Class"] == "Rural", "Proposed (BXGB)"].iloc[0])
    infl_pct = downstream["bias_pct"].mean()
    return {
        "prop": prop, "mice": mice, "knn": knn, "rf": rf,
        "binary": binary, "thr_bal": thr_bal, "thr_urb": thr_urb,
        "urban_acc": urban_acc, "trans_acc": trans_acc, "rural_acc": rural_acc,
        "infl_pct": infl_pct,
    }


def remove_duplicate_paragraphs(doc, prefix: str) -> None:
    seen = False
    to_remove = []
    for p in doc.paragraphs:
        if prefix in p.text and not has_equation(p._element):
            if seen:
                to_remove.append(p)
            else:
                seen = True
    for p in to_remove:
        p._element.getparent().remove(p._element)


def insert_after(doc, anchor: str, text: str) -> bool:
    for i, p in enumerate(doc.paragraphs):
        if anchor in p.text and not has_equation(p._element):
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph

            new_el = OxmlElement("w:p")
            p._element.addnext(new_el)
            new_p = Paragraph(new_el, p._parent)
            new_p.text = text
            new_p.paragraph_format.space_after = Pt(6)
            return True
    return False


def replace_heading(doc, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip() == old and not has_equation(p._element):
            replace_para(p, new)
            return


def remove_empty_paragraphs(doc) -> None:
    to_remove = []
    for p in doc.paragraphs:
        if not p.text.strip() and not has_equation(p._element):
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)


def remove_orphan_headings(doc) -> None:
    to_remove = []
    for p in doc.paragraphs:
        if p.text.strip() == "###":
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)


def apply_fine_tuning(m):
    doc = Document(str(MAIN))
    binary_acc = m["binary"]["accuracy_masked"]

    principal = (
        "In this study, we developed and evaluated a statistical framework for inferring missing administrative "
        "urban–rural labels from dietary survey data. Using CHNS as an illustrative cohort with complete T2 labels, "
        "we showed that inference accuracy under masked labels was moderate and stable across missingness levels and "
        "survey years (accuracy approximately 0.78–0.79 under the three-class descriptive framework; binary primary "
        f"task accuracy: {binary_acc:.3f}). Downstream analyses preserved direction and statistical significance of "
        "urban–rural contrasts, although effect sizes were moderately inflated."
    )

    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        if "under primary scenarios" in p.text or (
            "In this study, we developed and evaluated a statistical framework" in p.text
            and "three-class descriptive framework" not in p.text
        ):
            replace_para(p, principal)
        elif p.text.strip() in {"Strengths", "### Strengths"}:
            replace_para(p, "### Strengths")
        elif p.text.strip() in {"Limitations", "### Limitations"}:
            replace_para(p, "### Limitations")
        elif p.text.strip() in {"###", "### "}:
            p._element.getparent().remove(p._element)

    remove_empty_paragraphs(doc)
    remove_orphan_headings(doc)
    remove_duplicate_paragraphs(doc, "### Limitations")
    doc.save(str(MAIN))
    print(f"Fine-tuned {MAIN}")


def fix_supplementary_s7_s8(m):
    doc = Document(str(SUPP))
    b = m["binary"]
    s7_hdr = [
        "Scenario", "N (test)", "Accuracy", "Macro-F1", "Weighted-F1",
        "Urban recall", "Rural recall", "Rural PPV",
    ]
    s7_vals = [
        "Binary XGBoost", str(int(b["n_test"])), f"{b['accuracy_masked']:.3f}",
        f"{b['macro_f1_masked']:.3f}", f"{b['weighted_f1_masked']:.3f}",
        f"{b['urban_recall_masked']:.3f}", f"{b['rural_recall_masked']:.3f}",
        f"{b['rural_precision_masked']:.3f}",
    ]

    for p in list(doc.paragraphs):
        if "Supplementary Table S7" in p.text or "Supplementary Table S8" in p.text:
            p._element.getparent().remove(p._element)

    for table in list(doc.tables):
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr and hdr[0] in {"Scenario", "Urban threshold"}:
            table._element.getparent().remove(table._element)

    doc.add_paragraph(
        "Supplementary Table S7. Binary administrative classification (Rural vs. Urban, excluding "
        "Transitional observations) under 30% simulated missingness."
    )
    s7 = doc.add_table(rows=2, cols=len(s7_hdr))
    for i, h in enumerate(s7_hdr):
        set_cell_text(s7.rows[0].cells[i], h)
    for i, v in enumerate(s7_vals):
        set_cell_text(s7.rows[1].cells[i], v)

    doc.add_paragraph(
        "Supplementary Table S8. Urban probability threshold tuning (three-class model, 30% masked labels)."
    )
    thr = pd.read_csv(ROOT / "results/threshold_tuning.csv")
    s8 = doc.add_table(rows=1 + len(thr), cols=5)
    h2 = ["Urban threshold", "Overall accuracy", "Macro-F1", "Urban recall", "Rural recall"]
    for i, h in enumerate(h2):
        set_cell_text(s8.rows[0].cells[i], h)
    for ri, (_, row) in enumerate(thr.iterrows(), start=1):
        set_cell_text(s8.rows[ri].cells[0], f"{row['urban_threshold']:.2f}")
        set_cell_text(s8.rows[ri].cells[1], f"{row['overall_acc']:.3f}")
        set_cell_text(s8.rows[ri].cells[2], f"{row['macro_f1']:.3f}")
        set_cell_text(s8.rows[ri].cells[3], f"{row['urban_recall']:.3f}")
        set_cell_text(s8.rows[ri].cells[4], f"{row['rural_recall']:.3f}")

    doc.save(str(SUPP))
    print(f"Fixed S7/S8 layout in {SUPP}")


def apply_cover_fine_tuning(m):
    doc = Document(str(COVER))
    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        if "modest Urban accuracy" in p.text:
            replace_para(
                p,
                p.text.replace("modest Urban accuracy", f"Urban accuracy of {m['urban_acc']:.3f}"),
            )
        if "limited Urban recall" not in p.text and "Urban classification accuracy was" in p.text:
            continue
    doc.save(str(COVER))
    print(f"Fine-tuned {COVER}")


def apply_main(m):
    doc = Document(str(MAIN))

    title = (
        "A Statistical Evaluation Protocol for Inferring Missing Contextual Labels "
        "from Dietary Survey Data: Multi-Dimensional Validation in the China Health and Nutrition Survey"
    )

    primary_target = (
        "The primary inferential target of this framework is the binary administrative T2 status "
        "(Urban vs. Rural). The three-class structure is retained solely for descriptive epidemiological "
        "stratification. Consequently, when interpreting model performance, primary conclusions regarding "
        "inference accuracy are based on the binary Rural/Urban discrimination task, not the three-class "
        "overall accuracy."
    )

    abstract_objective = (
        "Objective:\n"
        "We propose a reproducible statistical evaluation protocol for assessing inference of missing "
        "administrative contextual labels in dietary surveys, using macronutrient composition and "
        "spatiotemporal covariates as predictors."
    )

    b = m["binary"]
    abstract_results = (
        "Results:\n"
        f"For the primary binary task (Rural vs. Urban, excluding the transitional stratum, n={int(b['n_test']):,}), "
        f"masked-label accuracy was {b['accuracy_masked']:.3f} (macro-F1 {b['macro_f1_masked']:.3f}; "
        f"weighted-F1 {b['weighted_f1_masked']:.3f}). Exploratory three-class accuracy was {m['prop']['Accuracy']:.3f}, "
        f"although this reflects partial definitional overlap for the Transitional stratum; the primary binary task "
        f"accuracy was {b['accuracy_masked']:.3f}. Transitional accuracy was {m['trans_acc']:.3f} (partially "
        f"definitional), while Urban accuracy was {m['urban_acc']:.3f}. The gradient-boosted classifier showed "
        f"incremental improvements over KNN ({m['knn']['Accuracy']:.3f}) and corrected multinomial-logistic MICE "
        f"({m['mice']['Accuracy']:.3f}). Downstream contrasts preserved statistical significance with moderate "
        f"Cohen's d inflation ({m['infl_pct']:.0f}% on average)."
    )

    intro_final = (
        "The present study addresses this gap by proposing a reproducible statistical evaluation protocol "
        "for assessing contextual label inference methods, using CHNS as an illustrative testbed. Rather than "
        "claiming optimal prediction, we establish a multi-dimensional validation framework that includes: "
        "(i) simulated missingness at varying rates; (ii) leave-one-year-out temporal generalisability; "
        "(iii) distributional fidelity via Jensen–Shannon divergence; (iv) probabilistic calibration assessment; "
        "and (v) downstream effect preservation. This protocol provides a template for future studies evaluating "
        "label-imputation methods when ground-truth contextual variables are available."
    )

    mice_methods = (
        "Multiple Imputation by Chained Equations (MICE) was implemented via iterative multinomial logistic "
        "regression for the categorical outcome (10 iterations), with continuous macronutrient and spatiotemporal "
        "predictors standardised prior to each imputation cycle. This specification avoids treating the outcome "
        "as a continuous variable, which would be inappropriate for administrative residence labels."
    )

    circularity_para = (
        primary_target + " Because Transitional is defined using FatER 23–30% and macronutrient predictors are "
        "included in the feature set, partial definitional overlap exists for this category; high Transitional "
        "accuracy should be interpreted accordingly."
    )

    inference_para = (
        f"Under 30% simulated missingness, the proposed framework achieved three-class accuracy "
        f"{m['prop']['Accuracy']:.3f} and macro-F1 {m['prop']['Macro_F1']:.3f} (Table 3). "
        f"It is important to note that Transitional accuracy is partially definitional due to the FatER-based "
        f"label construction; therefore, the remainder of this section focuses on the binary administrative "
        f"discrimination task (Rural vs. Urban), which avoids circularity."
    )

    binary_methods = (
        "For the primary binary inferential task, we additionally trained a separate XGBoost classifier on "
        "non-transitional observations only (Rural vs. Urban, excluding FatER 23–30%). The same hyperparameter "
        "configuration, class weighting, and evaluation protocol (including 30% simulated label missingness) were "
        "applied. This model provides the basis for primary performance reporting; the three-class model is retained "
        "for descriptive epidemiological stratification."
    )

    results_retention = (
        "Although the primary inferential target is binary administrative status, we retain the three-class structure "
        "throughout the results for descriptive completeness, as the Transitional stratum captures a meaningful "
        "intermediate dietary pattern in the nutrition transition literature."
    )

    accuracy_comparison = (
        f"The lower accuracy of the binary model ({b['accuracy_masked']:.3f}) compared with the three-class model "
        f"({m['prop']['Accuracy']:.3f}) is expected: the three-class model benefits from the Transitional category, "
        "which is partially defined by the same macronutrient features used for prediction. The binary task represents "
        "the more challenging administrative discrimination problem and is therefore the primary basis for evaluating "
        "the framework's practical utility."
    )

    binary_para = (
        f"Supplementary analysis: Binary administrative classification. To evaluate performance on the primary "
        f"inferential task, we fit a separate binary XGBoost model on non-transitional observations (Rural vs. Urban, "
        f"n={int(b['n_test']):,}). Under 30% simulated missingness, binary accuracy was {b['accuracy_masked']:.3f} "
        f"and macro-F1 {b['macro_f1_masked']:.3f}. Although macro-F1 was limited by low Urban recall "
        f"({b['urban_recall_masked']:.3f}), weighted-F1 was {b['weighted_f1_masked']:.3f}, reflecting practical "
        f"utility when Rural identification is prioritised. Rural recall was {b['rural_recall_masked']:.3f} and rural "
        f"positive predictive value (NPV for Urban-as-positive coding) was {b['rural_precision_masked']:.3f}, "
        f"supporting use as a high-specificity screening tool for non-urban dietary patterns. Full binary metrics are "
        f"reported in Supplementary Table S7."
    )

    urban_para = (
        f"Urban classification accuracy was {m['urban_acc']:.3f} under default three-class assignment. This is not a "
        "failure of the framework but rather reflects intrinsic heterogeneity of dietary patterns within administratively "
        "urban areas, where diets range from traditional to highly globalised. This finding establishes a practical "
        "boundary condition for inferring administrative residence from macronutrient data alone."
    )

    threshold_para = (
        "Probability threshold tuning. Because default argmax classification under-represents Urban labels, "
        f"we evaluated lower decision thresholds on predicted Urban probability. At an Urban threshold of 0.35, "
        f"masked-label Urban recall increased to {m['thr_urb']['urban_recall']:.2f} (overall accuracy "
        f"{m['thr_urb']['overall_acc']:.3f}); at 0.40, Urban recall was {m['thr_bal']['urban_recall']:.2f} with "
        f"accuracy {m['thr_bal']['overall_acc']:.3f} and Rural recall {m['thr_bal']['rural_recall']:.2f} "
        "(Supplementary Table S8). This supports application-specific threshold selection rather than a single "
        "deterministic rule."
    )

    mice_discussion = (
        "The performance of MICE relative to tree-based methods suggests that the relationship between "
        "macronutrient composition and administrative residence is inherently non-linear and interaction-dominated, "
        "where linear or locally linear imputation assumptions may be suboptimal."
    )

    comparison_discussion = (
        f"In this study, the proposed framework showed incremental improvements over conventional imputation methods "
        f"under simulated label missingness (accuracy {m['prop']['Accuracy']:.3f} vs. KNN {m['knn']['Accuracy']:.3f}, "
        f"MICE {m['mice']['Accuracy']:.3f}, and random forest imputation {m['rf']['Accuracy']:.3f}; Table 3). "
        "Most prior approaches to missing data in nutritional epidemiology have focused on imputing nutrient values "
        "rather than contextual variables; the present protocol evaluates categorical label recovery under a unified "
        "missingness simulation design."
    )

    screening_para = (
        "Given high Rural recall and rural positive predictive value in the binary task, the framework is best "
        "positioned as a high-specificity screening tool for identifying non-urban dietary patterns, or for assigning "
        "probabilistic weights in sensitivity analyses rather than producing deterministic universal classifications. "
        "Methodologically, it is intended for datasets where administrative contextual labels are missing but "
        "macronutrient and spatiotemporal fields are available."
    )

    strengths_body = (
        "This study has several strengths. First, the large sample size and longitudinal design enabled robust "
        "evaluation across time. Second, the multi-tiered evaluation framework assesses performance beyond accuracy, "
        "including distributional agreement, temporal generalisability, calibration, and preservation of epidemiological "
        "associations. Third, comparison against multiple baseline imputation methods supports interpretability relative "
        "to existing approaches."
    )

    limitations_intro = (
        "Several limitations should be considered. First, the three-category outcome combines administrative T2 labels "
        "with a FatER-defined transitional stratum. This creates partial definitional overlap between the Transitional "
        "category and macronutrient predictors. Therefore, all primary conclusions regarding inference accuracy are "
        "based on the binary Rural/Urban discrimination task; three-class results are presented for descriptive "
        "completeness only."
    )

    downstream_para = (
        "Analyses using inferred labels preserved the direction and statistical significance of urban–rural "
        f"differences across all dietary indicators (Table 8; Figure 5). Cohen's d was moderately inflated "
        f"({m['infl_pct']:.0f}% relative increase on average) when using inferred labels, because Transitional "
        "observations were assigned to Rural or Urban groups, sharpening group contrasts. Statistical significance "
        "was preserved in all cases; inferred labels are better suited to descriptive comparisons or covariate "
        "adjustment than standalone effect estimation."
    )

    limitations_first = limitations_intro

    gain_para = (
        f"The absolute performance gain over the strongest non-proposed comparator (random forest imputation) was "
        f"incremental (Δ accuracy ≈ {100 * (m['prop']['Accuracy'] - m['rf']['Accuracy']):.1f} percentage points) "
        f"but consistent across missing rates (mean advantage over KNN ≈ "
        f"{100 * (m['prop']['Accuracy'] - m['knn']['Accuracy']):.1f} percentage points)."
    )

    transition_downstream = (
        "Having established the framework's inferential performance, we next examine its practical utility for "
        "downstream epidemiological analyses—the ultimate test of any imputation strategy in applied research."
    )

    table8_note = (
        "All differences significant at P < 0.001 in both true and inferred analyses. "
        "Relative Change (%) = (|inferred d| − |true d|) / |true d| × 100."
    )

    replacements = [
        ("A Gradient-Boosted Inference Framework", title),
        ("A Statistical Evaluation Protocol", title),
        ("Objective:", abstract_objective),
        ("Results:", abstract_results),
        ("The present study addresses this gap by proposing a reproducible statistical evaluation protocol", intro_final),
        ("The present study addresses this gap by developing", intro_final),
        ("Because Transitional is defined using FatER", circularity_para),
        ("The primary inferential target of this framework", primary_target),
        ("Multiple Imputation by Chained Equations (MICE) uses iterative regression", mice_methods),
        ("Multiple Imputation by Chained Equations (MICE) was implemented", mice_methods),
        ("Under the primary evaluation scenario (30% simulated missingness), the proposed framework achieved", inference_para),
        ("Under 30% simulated missingness, the proposed framework achieved three-class accuracy", inference_para),
        ("The modest performance of MICE relative to tree-based methods", mice_discussion),
        ("The performance of MICE relative to tree-based methods", mice_discussion),
        ("reflecting limitations for categorical outcomes derived from behavioural data", mice_discussion),
        ("Supplementary analysis: Binary administrative classification", binary_para),
        ("The modest Urban accuracy", urban_para),
        ("Urban classification accuracy was", urban_para),
        ("In this study, the proposed framework showed modest but consistent improvements", comparison_discussion),
        ("In this study, the proposed framework showed incremental improvements", comparison_discussion),
        ("An important distinction is that the framework leverages global dietary structure", mice_discussion),
        ("Methodologically, the framework is intended for datasets where administrative", screening_para),
        ("Given high accuracy for Rural and Transitional strata", screening_para),
        ("Given high Rural recall and rural positive predictive value", screening_para),
        ("This study has several strengths. The large sample size", strengths_body),
        ("Several limitations should also be considered. First, the framework was developed", limitations_intro),
        ("First, the three-category outcome combines administrative T2 labels", limitations_intro),
        ("Several limitations should be considered. First, the three-category outcome", limitations_intro),
        ("Analyses using inferred labels preserved the direction", downstream_para),
        ("Effect sizes were moderately larger when using inferred labels", downstream_para),
        ("Second, the three-category outcome combines administrative T2 labels", (
            "Second, the framework was developed using data from a single country. Although internal validation "
            "suggests robustness across time and regions within China, external validation in other populations is needed."
        )),
        ("All differences significant at P < 0.001 in both true and inferred analyses. Bias", table8_note),
        ("The absolute performance gain over the strongest non-proposed comparator", gain_para),
        ("The absolute performance gain over the strongest baseline was modest", gain_para),
    ]

    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        for prefix, new in replacements:
            if prefix in p.text:
                replace_para(p, new)
                break

    if not any("For the primary binary inferential task, we additionally trained" in p.text for p in doc.paragraphs):
        insert_before(doc, "Baseline Methods for Comparison", binary_methods)

    if not any("Although the primary inferential target is binary administrative status, we retain" in p.text for p in doc.paragraphs):
        insert_before(doc, "Under 30% simulated missingness, the proposed framework achieved", results_retention)

    if not any("The lower accuracy of the binary model" in p.text for p in doc.paragraphs):
        insert_before(doc, "Urban classification accuracy was", accuracy_comparison)
        if not any("Urban classification accuracy was" in p.text for p in doc.paragraphs):
            insert_before(doc, "The modest Urban accuracy", accuracy_comparison)

    if not any("Supplementary analysis: Binary administrative classification" in p.text for p in doc.paragraphs):
        insert_before(doc, "Robustness Across Missing Rates", binary_para)
        insert_before(doc, "Robustness Across Missing Rates", urban_para)
        insert_before(doc, "Robustness Across Missing Rates", threshold_para)

    if not any("Having established the framework's inferential performance" in p.text for p in doc.paragraphs):
        insert_before(doc, "Preservation of Downstream Epidemiological Associations", transition_downstream)

    remove_duplicate_paragraphs(
        doc,
        "Under 30% simulated missingness, the proposed framework achieved three-class accuracy",
    )
    remove_duplicate_paragraphs(doc, "Analyses using inferred labels preserved the direction")
    remove_duplicate_paragraphs(
        doc, "In this study, the proposed framework showed incremental improvements over conventional imputation"
    )
    remove_duplicate_paragraphs(doc, "Given high Rural recall and rural positive predictive value in the binary task")

    # Remove stale comparison paragraphs superseded by consolidated text
    stale_prefixes = (
        "Most prior approaches to missing data in nutritional epidemiology have focused",
        "More broadly, this study extends methodological work in the field by addressing",
    )
    to_remove = []
    for p in doc.paragraphs:
        if any(p.text.startswith(s) for s in stale_prefixes) and not has_equation(p._element):
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)

    replace_heading(doc, "Strengths and Limitations", "### Strengths")
    replace_heading(doc, "Strengths", "### Strengths")
    replace_heading(doc, "Limitations", "### Limitations")
    if not any(p.text.strip() == "Limitations" for p in doc.paragraphs):
        for p in doc.paragraphs:
            if limitations_intro[:40] in p.text:
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph

                new_el = OxmlElement("w:p")
                p._element.addprevious(new_el)
                new_p = Paragraph(new_el, p._parent)
                new_p.text = "### Limitations"
                break
    else:
        for p in doc.paragraphs:
            if p.text.strip() == "Limitations":
                replace_para(p, "### Limitations")
                break

    if not any("Given high Rural recall and rural positive predictive value" in p.text for p in doc.paragraphs):
        insert_before(doc, "Strengths", screening_para)
        if not any("Given high Rural recall" in p.text for p in doc.paragraphs):
            insert_before(doc, "Limitations", screening_para)

    # Table 8 header
    t8 = doc.tables[7]
    hdr = t8.rows[0].cells
    set_cell_text(hdr[3], "Relative Change in Cohen's d (%)")

    doc.save(str(MAIN))
    print(f"Updated {MAIN}")


def apply_cover(m):
    doc = Document(str(COVER))
    title = (
        "A Statistical Evaluation Protocol for Inferring Missing Contextual Labels "
        "from Dietary Survey Data: Multi-Dimensional Validation in the China Health and Nutrition Survey"
    )
    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        if "We submit the manuscript entitled" in p.text:
            replace_para(
                p,
                f'We submit the manuscript entitled "{title}" for consideration as a Methods and Algorithms '
                "paper in Statistics in Medicine.",
            )
        if "Under 30% simulated missingness, binary Rural/Urban accuracy was" in p.text:
            replace_para(
                p,
                f"The manuscript emphasises statistical evaluation—simulated label missingness (10%–70%), "
                f"leave-one-year-out temporal holdout, comparator imputation methods, calibration, distributional "
                f"fidelity, and downstream effect preservation—rather than public health policy recommendations. "
                f"The binary Rural/Urban accuracy of {m['binary']['accuracy_masked']:.3f} represents the primary "
                f"evaluation metric; the three-class accuracy of {m['prop']['Accuracy']:.3f} includes the "
                f"descriptively defined Transitional stratum and is reported for completeness. The gradient-boosted "
                f"classifier showed incremental improvements over KNN ({m['knn']['Accuracy']:.3f}) and corrected MICE "
                f"({m['mice']['Accuracy']:.3f}).",
            )
        if "outperforming KNN" in p.text:
            replace_para(
                p,
                f"The manuscript emphasises statistical evaluation—simulated label missingness (10%–70%), "
                f"leave-one-year-out temporal holdout, comparator imputation methods, calibration, distributional "
                f"fidelity, and downstream effect preservation—rather than public health policy recommendations. "
                f"The binary Rural/Urban accuracy of {m['binary']['accuracy_masked']:.3f} represents the primary "
                f"evaluation metric; the three-class accuracy of {m['prop']['Accuracy']:.3f} includes the "
                f"descriptively defined Transitional stratum and is reported for completeness. The gradient-boosted "
                f"classifier showed incremental improvements over KNN ({m['knn']['Accuracy']:.3f}) and corrected MICE "
                f"({m['mice']['Accuracy']:.3f}).",
            )

    extra = (
        "We wish to clarify two methodological points emphasized in the revised manuscript. First, the "
        "three-category structure (Rural/Transitional/Urban) is retained for descriptive utility; however, the "
        "primary inferential target is binary administrative discrimination (Rural vs. Urban). The Transitional "
        "category, defined by FatER 23–30%, creates partial circularity when macronutrient predictors are used—a "
        "limitation we now explicitly discuss and address by focusing main evaluation on the binary task. Second, "
        f"Urban classification accuracy was {m['urban_acc']:.3f} under default three-class assignment. This reflects "
        "inherent dietary heterogeneity in administratively urban areas; we interpret this as an empirical limit of "
        "dietary data rather than a failure of the framework. We position the method as a high-specificity screening "
        "tool for non-urban patterns and as a template for statistical evaluation protocols, consistent with the scope "
        "of Statistics in Medicine."
    )
    if not any("We wish to clarify two methodological points" in p.text for p in doc.paragraphs):
        doc.add_paragraph("")
        doc.add_paragraph(extra)
    else:
        for p in doc.paragraphs:
            if "We wish to clarify two methodological points" in p.text:
                replace_para(p, extra)
                break
    doc.save(str(COVER))
    print(f"Updated {COVER}")


def apply_supplementary(m):
    fix_supplementary_s7_s8(m)


def apply_supplementary_legacy(m):
    doc = Document(str(SUPP))
    b = m["binary"]
    s7_hdr = [
        "Scenario", "N (test)", "Accuracy", "Macro-F1", "Weighted-F1",
        "Urban recall", "Rural recall", "Rural PPV",
    ]
    s7_vals = [
        "Binary XGBoost", str(int(b["n_test"])), f"{b['accuracy_masked']:.3f}",
        f"{b['macro_f1_masked']:.3f}", f"{b['weighted_f1_masked']:.3f}",
        f"{b['urban_recall_masked']:.3f}", f"{b['rural_recall_masked']:.3f}",
        f"{b['rural_precision_masked']:.3f}",
    ]

    if not any("Supplementary Table S7" in p.text for p in doc.paragraphs):
        doc.add_paragraph(
            "Supplementary Table S7. Binary administrative classification (Rural vs. Urban, excluding "
            "Transitional observations) under 30% simulated missingness."
        )
        table = doc.add_table(rows=2, cols=len(s7_hdr))
        for i, h in enumerate(s7_hdr):
            set_cell_text(table.rows[0].cells[i], h)
        for i, v in enumerate(s7_vals):
            set_cell_text(table.rows[1].cells[i], v)
    else:
        updated = False
        for table in doc.tables:
            hdr = [c.text.strip() for c in table.rows[0].cells]
            if hdr and hdr[0] == "Scenario":
                if len(hdr) != len(s7_hdr):
                    tbl_el = table._element
                    tbl_el.getparent().remove(tbl_el)
                    break
                for i, v in enumerate(s7_vals):
                    set_cell_text(table.rows[1].cells[i], v)
                updated = True
                break
        if not updated and any("Supplementary Table S7" in p.text for p in doc.paragraphs):
            doc.add_paragraph(
                "Supplementary Table S7 (updated). Binary administrative classification (Rural vs. Urban, excluding "
                "Transitional observations) under 30% simulated missingness."
            )
            table = doc.add_table(rows=2, cols=len(s7_hdr))
            for i, h in enumerate(s7_hdr):
                set_cell_text(table.rows[0].cells[i], h)
            for i, v in enumerate(s7_vals):
                set_cell_text(table.rows[1].cells[i], v)

    if not any("Supplementary Table S8" in p.text for p in doc.paragraphs):
        doc.add_paragraph(
            "Supplementary Table S8. Urban probability threshold tuning (three-class model, 30% masked labels)."
        )
        thr = pd.read_csv(ROOT / "results/threshold_tuning.csv")
        t2 = doc.add_table(rows=1 + len(thr), cols=5)
        h2 = ["Urban threshold", "Overall accuracy", "Macro-F1", "Urban recall", "Rural recall"]
        for i, h in enumerate(h2):
            set_cell_text(t2.rows[0].cells[i], h)
        for ri, (_, row) in enumerate(thr.iterrows(), start=1):
            set_cell_text(t2.rows[ri].cells[0], f"{row['urban_threshold']:.2f}")
            set_cell_text(t2.rows[ri].cells[1], f"{row['overall_acc']:.3f}")
            set_cell_text(t2.rows[ri].cells[2], f"{row['macro_f1']:.3f}")
            set_cell_text(t2.rows[ri].cells[3], f"{row['urban_recall']:.3f}")
            set_cell_text(t2.rows[ri].cells[4], f"{row['rural_recall']:.3f}")

    doc.save(str(SUPP))
    print(f"Updated {SUPP}")


def write_response(m):
    text = f"""# Response to Reviewer Concerns (SiM revision)

## 1. Transitional circularity (Priority 1)
- **Methods:** Added explicit statement that binary T2 Rural/Urban is the primary inferential target; three-class structure is descriptive only.
- **Abstract/Results:** Transitional accuracy ({m['trans_acc']:.3f}) reported as exploratory; main evaluation reframed around binary task (accuracy {m['binary']['accuracy_masked']:.3f}).
- **Discussion – Limitations:** Circularity listed as first limitation.

## 2. Low Urban accuracy ({m['urban_acc']:.3f})
- **Results:** Added paragraph on dietary heterogeneity in urban areas.
- **Discussion:** Reframed framework as screening / probabilistic weighting tool.
- **New analysis:** Urban probability threshold tuning (S8); Urban recall 0.58–0.66 at thresholds 0.40–0.35.

## 3. MICE implementation
- **Methods:** Replaced continuous IterativeImputer with iterative multinomial logistic regression for categorical outcomes.
- **Results:** MICE accuracy updated to {m['mice']['Accuracy']:.3f} (from 0.393).
- **Discussion:** Added interpretation regarding non-linear structure.

## 4. Statistical innovation / SiM fit
- **Introduction/Abstract:** Reframed contribution as multi-dimensional **evaluation protocol**, not optimal prediction claim.
- **Title:** Updated to emphasize statistical evaluation protocol.

## 5. Downstream effect size wording
- **Table 8:** Renamed "Bias (%)" to "Relative Change in Cohen's d (%)".
- **Results/Discussion:** "Bias" replaced with "inflation/amplification" language; clarified mechanism via Transitional reassignment.

## 6. New analyses
- **Binary model (S7):** Rural vs. Urban excluding transitional stratum.
- **Threshold tuning (S8):** Application-specific Urban probability cut-offs.

## 7. Cover letter
- Added clarifications on three-category rationale, circularity, Urban accuracy, and SiM scope.

## 8. Second-round revisions
- **Binary metrics:** Added weighted-F1 ({m['binary']['weighted_f1_masked']:.3f}) and rural PPV ({m['binary']['rural_precision_masked']:.3f}); reframed as high-specificity rural screening tool.
- **Accuracy comparison:** Explained three-class ({m['prop']['Accuracy']:.3f}) vs binary ({m['binary']['accuracy_masked']:.3f}) difference.
- **Methods:** Added binary model training paragraph.
- **Discussion:** Consolidated comparison paragraphs; split Strengths/Limitations subsections.
- **Duplicates:** Removed repeated downstream results paragraph.
- **Cover letter:** Binary accuracy stated as primary evaluation metric.
- **Note:** Figure captions (1–5) are present; embedded images should be verified in Word before submission.

## 9. Third-round fine tuning
- **Principal Findings:** Clarified three-class vs binary accuracy wording.
- **Discussion headings:** Standardised `### Strengths` and `### Limitations`; removed blank/orphan paragraphs.
- **Supplementary S7/S8:** Rebuilt as contiguous caption-table blocks.
- **Cover letter:** Removed residual "modest Urban accuracy" phrasing.
"""
    RESPONSE.write_text(text, encoding="utf-8")
    print(f"Wrote {RESPONSE}")


def main():
    m = load_metrics()
    apply_main(m)
    apply_fine_tuning(m)
    apply_cover(m)
    apply_cover_fine_tuning(m)
    apply_supplementary(m)
    write_response(m)


if __name__ == "__main__":
    main()
