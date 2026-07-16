#!/usr/bin/env python3
"""Apply SiM major-revision round: selection bias, abstract order, sensitivities, LOYO decline."""
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MAIN = ROOT / "public" / "main.docx"
SUPP = ROOT / "public" / "Supplementary Material.docx"
RESPONSE = ROOT / "public" / "RESPONSE_TO_REVIEWERS.md"


def has_eq(el) -> bool:
    for n in el.iter():
        tag = n.tag.split("}")[-1] if "}" in n.tag else n.tag
        if tag in ("oMath", "oMathPara"):
            return True
    return False


def set_text(p, text: str) -> None:
    if not has_eq(p._element):
        p.text = text


def insert_before(doc, anchor: str, text: str) -> bool:
    for p in doc.paragraphs:
        if anchor in p.text and not has_eq(p._element):
            np_ = p.insert_paragraph_before(text)
            np_.paragraph_format.space_after = Pt(6)
            return True
    return False


def dedupe(doc, prefix: str) -> None:
    seen = False
    rm = []
    for p in doc.paragraphs:
        if prefix in p.text and not has_eq(p._element):
            if seen:
                rm.append(p)
            else:
                seen = True
    for p in rm:
        p._element.getparent().remove(p._element)


def load():
    binary = pd.read_csv(ROOT / "results/binary_classification.csv").iloc[0]
    noise = pd.read_csv(ROOT / "results/binary_noise_injection_sensitivity.csv")
    mar = pd.read_csv(ROOT / "results/mar_beta_sensitivity.csv")
    hp = pd.read_csv(ROOT / "results/hyperparameter_perturbation_summary.csv").iloc[0]
    loyo = pd.read_csv(ROOT / "results/leave_one_year_out.csv")
    class_acc = pd.read_csv(ROOT / "results/enhanced_baseline_class_acc.csv")
    baseline = pd.read_csv(ROOT / "results/enhanced_baseline_comparison.csv")
    adj = pd.read_csv(ROOT / "results/downstream_adjusted_regression.csv")
    prop = baseline.loc[baseline["Method"] == "Proposed (BXGB)"].iloc[0]
    trans_acc = float(class_acc.loc[class_acc["Class"] == "Transitional", "Proposed (BXGB)"].iloc[0])
    urban_acc = float(class_acc.loc[class_acc["Class"] == "Urban", "Proposed (BXGB)"].iloc[0])

    primary = noise.loc[noise["scenario"] == "primary_exclude_transitional"].iloc[0]
    by_t2 = noise.loc[noise["scenario"] == "collapse_transitional_by_T2"].iloc[0]
    rand = noise.loc[noise["scenario"] == "collapse_transitional_random50"].iloc[0]

    loyo_1993 = float(loyo.loc[loyo["Test_Year"] == 1993, "Imputation_Accuracy"].iloc[0])
    loyo_2011 = float(loyo.loc[loyo["Test_Year"] == 2011, "Imputation_Accuracy"].iloc[0])
    fat = adj.loc[adj["outcome"] == "FatER"].iloc[0]

    return dict(
        binary=binary, prop=prop, trans_acc=trans_acc, urban_acc=urban_acc,
        primary=primary, by_t2=by_t2, rand=rand, mar=mar, hp=hp,
        loyo_1993=loyo_1993, loyo_2011=loyo_2011, fat=fat, adj=adj,
    )


def apply_main(m):
    doc = Document(str(MAIN))

    abstract_results = (
        "Results:\n"
        f"For the primary binary task (Rural vs Urban, excluding the transitional FatER stratum, "
        f"n={int(m['binary']['n_test']):,}), masked-label accuracy was {m['binary']['accuracy_masked']:.3f} "
        f"(macro-F1 {m['binary']['macro_f1_masked']:.3f}; weighted-F1 {m['binary']['weighted_f1_masked']:.3f}). "
        f"For descriptive completeness only, exploratory three-class accuracy was {m['prop']['Accuracy']:.3f}; "
        f"Transitional accuracy ({m['trans_acc']:.3f}) is substantially definitional because that stratum is "
        f"thresholded on FatER, a predictor, whereas Urban accuracy was {m['urban_acc']:.3f}. "
        f"Downstream contrasts preserved significance with moderate Cohen's d inflation "
        f"(~25% unadjusted; adjusted Urban FatER coefficient change {m['fat']['coef_rel_change_pct']:.0f}% "
        f"after Year and Province controls)."
    )

    sample_note = (
        "Analytic covariates. The CHNS diet-file extract used here provides macronutrients, survey year, "
        "province, and administrative T2, but does not include age or sex. All modelling and adjusted "
        "downstream regressions therefore condition only on available fields (Year and Province indicators "
        "in Supplementary Table S10). Conclusions about covariate-adjusted effect preservation should not "
        "be extrapolated to age–sex–adjusted epidemiological models without additional linked demographic data."
    )

    selection_bias = (
        f"Selection implications of the binary primary task. Excluding Transitional observations "
        f"(FatER 23–30%) removes the dietary overlap region between administrative Rural and Urban profiles and "
        f"can inflate apparent separability. Under 30% MCAR masking, primary binary accuracy was "
        f"{m['primary']['masked_accuracy']:.3f} (n={int(m['primary']['n_test']):,}). When Transitional labels "
        f"were retained but collapsed to binary using true T2, masked accuracy remained similar "
        f"({m['by_t2']['masked_accuracy']:.3f}; n={int(m['by_t2']['n_test']):,}). Random 50/50 assignment of "
        f"Transitional labels reduced accuracy to {m['rand']['masked_accuracy']:.3f} (Supplementary Table S11). "
        "Thus the headline binary figure should be interpreted as performance on the non-overlap administrative "
        "stratum; external validity to populations that retain transitional diets requires the T2-collapse or "
        "noise-injection benchmarks, not the exclusion-only estimate alone."
    )

    mar_sens = (
        f"MAR propensity slope sensitivity. Holding the realised missing rate near 30%, we varied "
        f"β₁ ∈ {{1.0, 1.5, 2.0, 2.5}}. Proposed three-class masked accuracy declined monotonically from "
        f"{float(m['mar'].loc[m['mar']['beta1']==1.0,'accuracy'].iloc[0]):.3f} (β₁=1.0) to "
        f"{float(m['mar'].loc[m['mar']['beta1']==2.5,'accuracy'].iloc[0]):.3f} (β₁=2.5) "
        f"(Supplementary Table S12). The primary choice β₁=2.0 (accuracy "
        f"{float(m['mar'].loc[m['mar']['beta1']==2.0,'accuracy'].iloc[0]):.3f}) is therefore a mid-range "
        "stress test of FatER-dependent missingness rather than a uniquely optimised value; empirical "
        "calibration of β₁ to external incomplete surveys remains future work when such incomplete labels exist."
    )

    hp_sens = (
        f"Hyperparameter stability. Perturbing max_depth ∈ {{3,4,5}}, learning_rate ∈ {{0.04,0.05,0.06}}, "
        f"and subsample ∈ {{0.72,0.90,1.00}} (28 settings) yielded mean masked accuracy {m['hp']['acc_mean']:.3f} "
        f"(SD {m['hp']['acc_sd']:.3f}); the largest absolute deviation from the baseline configuration was "
        f"{m['hp']['max_abs_delta_pp']:.2f} percentage points (Supplementary Table S13). Within this neighbourhood, "
        "performance is stable for protocol reuse."
    )

    loyo_disc = (
        f"Temporal transportability is not flat. Leave-one-year-out accuracy peaked near 1993 "
        f"({m['loyo_1993']:.3f}) and declined to {m['loyo_2011']:.3f} when 2011 was held out "
        f"(Table 6)—about a {100*(m['loyo_1993']-m['loyo_2011']):.1f} percentage-point drop—despite an "
        "acceptable overall mean. This pattern indicates mild temporal drift: dietary–residence associations "
        "evolve, so applications to newer waves should retrain on recent labelled cohorts or use "
        "time-adaptive reweighting rather than freezing a single historical model."
    )

    table2_note = (
        "Note on Table 2. Under the T2-plus-overlay definition, Rural excludes FatER ∈ [23%, 30%] but still "
        "includes administrative rural observations with FatER > 30% (urbanised diets) and FatER < 23%. "
        "Consequently the Rural FatER mean can sit near 0.25 without any Rural observation lying inside the "
        "transitional band; the mean reflects a mixture of traditional and urbanised rural diets, not "
        "definitional leakage of the transitional stratum."
    )

    spatial_clarify = (
        "Spatially structured missingness uses observation-level Bernoulli draws with province-specific "
        "probabilities (0.50 in Beijing, Shanghai, and Chongqing; 0.20 elsewhere). Entire provinces are not "
        "deleted; within each province, masking remains random at the observation level."
    )

    downstream_lim = (
        "Downstream effect preservation in the main text emphasises unadjusted Cohen's d (Table 8) for "
        "transparency with prior descriptive contrasts. Covariate-adjusted Urban coefficients "
        "(Year + Province; Supplementary Table S10) showed similar directional preservation with "
        f"relative coefficient change of {m['fat']['coef_rel_change_pct']:.0f}% for FatER. Because age and sex "
        "are unavailable in this extract, these adjusted results do not replace age–sex–adjusted "
        "epidemiological practice; they demonstrate protocol behaviour under the richest adjustment set "
        "available here."
    )

    # Replace / terminology
    for p in doc.paragraphs:
        if has_eq(p._element):
            continue
        t = p.text
        if t.startswith("Results:\n") or (t.startswith("Results:") and "primary binary" in t):
            set_text(p, abstract_results)
        elif "partially definitional" in t:
            set_text(p, t.replace("partially definitional", "substantially definitional"))
        elif "partial definitional" in t:
            set_text(p, t.replace("partial definitional", "substantial definitional"))

    # Insert new paragraphs if missing
    if not any("Analytic covariates. The CHNS diet-file extract" in p.text for p in doc.paragraphs):
        insert_before(doc, "Outcome Definition", sample_note)
        if not any("Analytic covariates" in p.text for p in doc.paragraphs):
            insert_before(doc, "Rural:", sample_note)

    if not any("Selection implications of the binary primary task" in p.text for p in doc.paragraphs):
        insert_before(doc, "Threshold fragility of the three-class", selection_bias)

    if not any("MAR propensity slope sensitivity" in p.text for p in doc.paragraphs):
        insert_before(doc, "Threshold fragility of the three-class", mar_sens)

    if not any("Hyperparameter stability" in p.text for p in doc.paragraphs):
        insert_before(doc, "Supplementary Analyses", hp_sens)

    if not any("Temporal transportability is not flat" in p.text for p in doc.paragraphs):
        insert_before(doc, "Protocol application guide", loyo_disc)

    if not any("Note on Table 2" in p.text for p in doc.paragraphs):
        insert_before(doc, "Under 30% simulated missingness, the proposed framework achieved three-class", table2_note)

    if not any("observation-level Bernoulli draws with province-specific" in p.text for p in doc.paragraphs):
        # strengthen spatial sentence in missingness section
        for p in doc.paragraphs:
            if "Spatially structured missingness" in p.text and "Beijing" in p.text:
                # append clarification if not present
                if "Entire provinces are not deleted" not in p.text:
                    set_text(p, p.text.rstrip() + " " + spatial_clarify)
                break

    if not any("age–sex–adjusted epidemiological practice" in p.text for p in doc.paragraphs):
        insert_before(doc, "Protocol application guide", downstream_lim)

    # Move retraining emphasis to Future Directions
    future_retrain = (
        "External application to newer survey waves requires retraining on recent labelled cohorts or "
        "time-adaptive reweighting, using at minimum updated macronutrient, year, province, and "
        "administrative T2 fields (and ideally age and sex when linkable). Frozen historical models should "
        "not be assumed transportable indefinitely given the LOYO decline toward 2011."
    )
    if not any("External application to newer survey waves requires retraining" in p.text for p in doc.paragraphs):
        insert_before(doc, "Finally, evaluating the impact of this approach", future_retrain)
    else:
        for p in doc.paragraphs:
            if "External application to newer survey waves requires retraining" in p.text:
                set_text(p, future_retrain)
                break

    # Table footnotes: append note near Table 3/4 captions if present
    for p in doc.paragraphs:
        if has_eq(p._element):
            continue
        if p.text.strip().startswith("Table 3") or p.text.strip().startswith("Table 4"):
            if "substantially definitional" not in p.text:
                set_text(
                    p,
                    p.text.rstrip()
                    + " Note: Transitional accuracy is substantially definitional (FatER-threshold overlay).",
                )

    # Deduplicate
    for pref in [
        "Selection implications of the binary primary task",
        "MAR propensity slope sensitivity",
        "Hyperparameter stability",
        "Temporal transportability is not flat",
        "Analytic covariates. The CHNS diet-file extract",
        "Note on Table 2",
        "Downstream effect preservation in the main text emphasises",
        "External application to newer survey waves requires retraining",
    ]:
        dedupe(doc, pref)

    doc.save(str(MAIN))
    print(f"Updated {MAIN}")


def append_supp_tables(m):
    doc = Document(str(SUPP))

    def add_caption(text):
        p = doc.add_paragraph(text)
        for r in p.runs:
            r.bold = True

    if not any("Supplementary Table S11" in p.text for p in doc.paragraphs):
        add_caption(
            "Supplementary Table S11. Binary-task selection-bias sensitivity: Transitional exclusion "
            "versus collapse-by-T2 versus random 50/50 assignment (30% MCAR masking)."
        )
        doc.add_paragraph(
            "Primary binary excludes FatER 23–30%. Collapse-by-T2 retains all observations and assigns "
            "Transitional to their administrative T2. Random 50/50 injects label noise in the overlap region."
        )
        noise = pd.read_csv(ROOT / "results/binary_noise_injection_sensitivity.csv")
        t = doc.add_table(rows=1 + len(noise), cols=6)
        hdr = ["Scenario", "N (test)", "Masked accuracy", "Masked F1", "Holdout accuracy", "Urban share"]
        for j, h in enumerate(hdr):
            t.rows[0].cells[j].paragraphs[0].text = h
        labels = {
            "primary_exclude_transitional": "Exclude Transitional (primary)",
            "collapse_transitional_by_T2": "Collapse Transitional by T2",
            "collapse_transitional_random50": "Random 50/50 for Transitional",
        }
        for i, (_, r) in enumerate(noise.iterrows(), start=1):
            vals = [
                labels.get(r["scenario"], r["scenario"]),
                str(int(r["n_test"])),
                f"{r['masked_accuracy']:.3f}",
                f"{r['masked_f1']:.3f}",
                f"{r['holdout_accuracy']:.3f}",
                f"{r['urban_share']:.3f}",
            ]
            for j, v in enumerate(vals):
                t.rows[i].cells[j].paragraphs[0].text = v

    if not any("Supplementary Table S12" in p.text for p in doc.paragraphs):
        add_caption(
            "Supplementary Table S12. MAR propensity slope sensitivity (target realised missingness ≈ 30%)."
        )
        mar = m["mar"]
        t = doc.add_table(rows=1 + len(mar), cols=5)
        for j, h in enumerate(["β₁", "β₀", "Realised rate", "Accuracy", "Macro-F1"]):
            t.rows[0].cells[j].paragraphs[0].text = h
        for i, (_, r) in enumerate(mar.iterrows(), start=1):
            for j, v in enumerate([
                f"{r['beta1']:.1f}", f"{r['beta0']:.3f}", f"{r['realized_rate']:.3f}",
                f"{r['accuracy']:.3f}", f"{r['macro_f1']:.3f}",
            ]):
                t.rows[i].cells[j].paragraphs[0].text = v

    if not any("Supplementary Table S13" in p.text for p in doc.paragraphs):
        add_caption(
            "Supplementary Table S13. Hyperparameter perturbation summary (± neighbourhood of Table S6)."
        )
        hp = m["hp"]
        doc.add_paragraph(
            f"Across {int(hp['n_settings'])} settings of max_depth, learning_rate, and subsample, "
            f"masked accuracy mean={hp['acc_mean']:.3f}, SD={hp['acc_sd']:.3f}, "
            f"range [{hp['acc_min']:.3f}, {hp['acc_max']:.3f}]; "
            f"max |Δ| from baseline = {hp['max_abs_delta_pp']:.2f} percentage points. "
            "Full grid in results/hyperparameter_perturbation.csv."
        )

    # Fig S1 / wording: substantially definitional
    for p in doc.paragraphs:
        if "partially definitional" in p.text or "partial definitional" in p.text:
            set_text(
                p,
                p.text.replace("partially definitional", "substantially definitional").replace(
                    "partial definitional", "substantial definitional"
                ),
            )

    doc.save(str(SUPP))
    print(f"Updated {SUPP}")


def write_response(m):
    text = f"""# Response to SiM Major Revision (second round)

## Concern 1 — Binary-task selection bias
- Added explicit discussion of excluding the FatER overlap stratum.
- New sensitivity (Supplementary Table S11): primary exclude Transitional masked acc={m['primary']['masked_accuracy']:.3f};
  collapse-by-T2={m['by_t2']['masked_accuracy']:.3f}; random 50/50={m['rand']['masked_accuracy']:.3f}.
- Script: `binary_noise_injection_sensitivity.py`.

## Concern 2 — Three-class presentation
- Abstract now leads with binary {m['binary']['accuracy_masked']:.3f}; three-class {m['prop']['Accuracy']:.3f} marked descriptive.
- Wording upgraded to **substantially definitional**; Table 3/4 captions footnoted.

## Concern 3 — MAR β₁ choice
- Sensitivity β₁∈{{1.0,1.5,2.0,2.5}} (Supplementary Table S12): accuracy {m['mar']['accuracy'].min():.3f}–{m['mar']['accuracy'].max():.3f}.
- Script: `mar_beta_sensitivity.py`.

## Concern 4 — Adjusted downstream
- Methods now pre-disclose age/sex unavailable; S10 Year+Province adjustment highlighted in Discussion as limited but available adjustment set.
- Unadjusted Table 8 retained as descriptive; adjusted FatER Δcoef≈{m['fat']['coef_rel_change_pct']:.0f}%.

## Concern 5 — Hyperparameter sensitivity
- ± neighbourhood grid (Supplementary Table S13): SD={m['hp']['acc_sd']:.3f}; max |Δ|={m['hp']['max_abs_delta_pp']:.2f} pp.
- Script: `hyperparameter_perturbation.py`.

## Concern 6 — LOYO decline
- Discussion now states decline from {m['loyo_1993']:.3f} (1993) to {m['loyo_2011']:.3f} (2011) and recommends retraining / time-adaptive weighting (Future Directions).

## Minor
- Spatial masking clarified as observation-level Bernoulli with province-specific rates.
- Table 2 Rural FatER mean≈0.25 explained (admin rural with FatER>30% allowed).
- Retraining moved/expanded under Future Directions.
"""
    RESPONSE.write_text(text, encoding="utf-8")
    print(f"Wrote {RESPONSE}")


def main():
    m = load()
    apply_main(m)
    append_supp_tables(m)
    write_response(m)


if __name__ == "__main__":
    main()
