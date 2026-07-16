#!/usr/bin/env python3
"""Apply SiM major-revision narrative upgrades to main.docx + Supplementary."""
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MAIN = ROOT / "public" / "main.docx"
SUPP = ROOT / "public" / "Supplementary Material.docx"
COVER = ROOT / "public" / "cover_letter_SiM.docx"
RESPONSE = ROOT / "public" / "RESPONSE_TO_REVIEWERS.md"


def has_equation(element) -> bool:
    for node in element.iter():
        tag = node.tag.split("}")[-1] if "}" in node.tag else node.tag
        if tag in ("oMath", "oMathPara"):
            return True
    return False


def replace_para(p, text: str) -> None:
    if has_equation(p._element):
        return
    p.text = text


def insert_before(doc, anchor: str, text: str) -> bool:
    for p in doc.paragraphs:
        if anchor in p.text and not has_equation(p._element):
            new_p = p.insert_paragraph_before(text)
            new_p.paragraph_format.space_after = Pt(6)
            return True
    return False


def insert_heading_before(doc, anchor: str, heading: str) -> bool:
    for p in doc.paragraphs:
        if anchor in p.text and not has_equation(p._element):
            new_p = p.insert_paragraph_before(heading)
            new_p.paragraph_format.space_before = Pt(10)
            new_p.paragraph_format.space_after = Pt(4)
            for r in new_p.runs:
                r.bold = True
            return True
    return False


def remove_duplicates(doc, prefix: str) -> None:
    seen = False
    to_remove = []
    for p in doc.paragraphs:
        if prefix in p.text and not has_equation(p._element):
            if seen:
                to_remove.append(p)
            else:
                seen = True
    for p in to_remove:
        p._element.getparent().remove(p._element)


def load_metrics():
    mech = pd.read_csv(ROOT / "results/missingness_mechanism_params.csv")
    miss = pd.read_csv(ROOT / "results/missingness_simulation.csv")
    adj = pd.read_csv(ROOT / "results/downstream_adjusted_regression.csv")
    band = pd.read_csv(ROOT / "results/threshold_sensitivity.csv")
    binary = pd.read_csv(ROOT / "results/binary_classification.csv").iloc[0]
    baseline = pd.read_csv(ROOT / "results/enhanced_baseline_comparison.csv")
    prop = baseline.loc[baseline["Method"] == "Proposed (BXGB)"].iloc[0]

    mar = mech.loc[mech["mechanism"] == "MAR"].iloc[0]
    mcar_acc = float(miss.loc[(miss["Scenario"] == "MCAR") & (miss["Method"] == "Proposed"), "Accuracy"].iloc[0])
    mar_acc = float(miss.loc[(miss["Scenario"] == "MAR (FatER)") & (miss["Method"] == "Proposed"), "Accuracy"].iloc[0])
    spat_acc = float(miss.loc[(miss["Scenario"] == "Spatial") & (miss["Method"] == "Proposed"), "Accuracy"].iloc[0])

    band_default = band.loc[(band["trans_low"] == 0.23) & (band["trans_high"] == 0.30)].iloc[0]
    band_wide = band.loc[(band["trans_low"] == 0.25) & (band["trans_high"] == 0.32)].iloc[0]

    fat = adj.loc[adj["outcome"] == "FatER"].iloc[0]
    return {
        "mar": mar, "mcar_acc": mcar_acc, "mar_acc": mar_acc, "spat_acc": spat_acc,
        "band_default": band_default, "band_wide": band_wide, "fat": fat,
        "binary": binary, "prop": prop, "adj": adj,
    }


def apply_main(m):
    doc = Document(str(MAIN))

    pillars = (
        "We formalise a reusable statistical evaluation protocol for contextual-label imputation with three pillars. "
        "Pillar 1 — Simulated missingness for internal validation: when a survey provides complete labels for a "
        "reference cohort, labels can be masked under MCAR/MAR/spatially structured mechanisms to stress-test "
        "imputation methods before deployment. Pillar 2 — Separate predictive accuracy from downstream inferential "
        "preservation: high classification accuracy does not guarantee that exposure contrasts, regression "
        "coefficients, or standard errors remain usable for epidemiology. Pillar 3 — Spatiotemporal generalisability: "
        "leave-one-year-out and leave-one-province-out holdouts quantify transportability across survey waves and "
        "regions. These pillars matter because missing contextual labels threaten both prediction and causal/"
        "associational analyses in nutritional epidemiology; a protocol that only optimises accuracy can mislead."
    )

    missing_heading_anchor = "Primary evaluations masked contextual labels completely at random"
    missing_section = (
        "MCAR. On the held-out test set, each observation's contextual label is masked independently with "
        "probability π = 0.30 (primary scenario; robustness also uses π ∈ {0.10,…,0.70}). Masking is independent of "
        "covariates and of the true label. Predictors remain fully observed.\n\n"
        f"MAR. We generate missingness depending on observed FatER via a logistic propensity model: "
        f"logit{{P(R=1 | FatER)}} = β₀ + β₁ z(FatER), where z(FatER) is the standardised FatER on the test set, "
        f"β₁ = {m['mar']['beta1']:.1f}, and β₀ is calibrated so that E[P(R=1)] ≈ 0.30 "
        f"(fitted β₀ = {m['mar']['beta0']:.3f}; realised rate {m['mar']['realized_rate']:.3f}). "
        "Larger FatER therefore increases the probability of label missingness, approximating settings where "
        "more urbanised dietary profiles are differentially incomplete.\n\n"
        "Spatially structured missingness. Labels are masked at the observation level with province-specific "
        "rates: 0.50 in Beijing, Shanghai, and Chongqing and 0.20 elsewhere (not whole-province deletion). "
        "This sensitivity check is reported alongside MCAR/MAR; primary conclusions rely on MCAR with MAR as "
        "the main departure-from-randomness stress test."
    )

    cal_rationale = (
        "Calibration is evaluated because well-calibrated class probabilities are not merely a goodness-of-fit "
        "summary: they can be used as inverse-probability weights or as draws in multiple imputation for "
        "downstream sensitivity analyses when labels are missing. We therefore report Brier scores and expected "
        "calibration error (ECE) and interpret calibration as a prerequisite for probability-based reuse of "
        "imputed labels, not only as a classifier diagnostic."
    )

    js_rationale = (
        "Jensen–Shannon divergence summarises agreement between observed and imputed category marginals. "
        "JS → 0 is necessary but not sufficient for faithful recovery: marginal balance can hold while "
        "year–region–category joints remain distorted (e.g., the urban share in a specific wave). We therefore "
        "treat low JS as a screening criterion and rely on temporal/spatial holdouts and downstream regressions "
        "to probe joint structure beyond margins."
    )

    s9_results = (
        f"Threshold fragility of the three-class descriptive labels. Varying the transitional FatER band "
        f"changed three-class accuracy sharply: under the default band [0.23, 0.30] accuracy was "
        f"{m['band_default']['accuracy']:.3f}, whereas widening to [0.25, 0.32] reduced accuracy to "
        f"{m['band_wide']['accuracy']:.3f} (Supplementary Table S9). This instability shows that hard thresholds "
        "on a predictor used to define an outcome render multi-class accuracy an unreliable headline metric. "
        f"Accordingly, primary conclusions rest on the binary T2 Rural/Urban task "
        f"(masked accuracy {m['binary']['accuracy_masked']:.3f}; Supplementary Table S7)."
    )

    adj_results = (
        f"Covariate-adjusted downstream preservation. Beyond unadjusted Cohen's d, we fit OLS models of each "
        f"macronutrient outcome on Urban versus Rural status, adjusting for survey year and province indicators "
        f"(age and sex are unavailable in this diet-file extract). Under 30% MCAR masking, the Urban coefficient "
        f"for FatER changed from {m['fat']['true_coef']:.4f} (SE {m['fat']['true_se']:.4f}) with true labels to "
        f"{m['fat']['imp_coef']:.4f} (SE {m['fat']['imp_se']:.4f}) with imputed labels "
        f"(relative change {m['fat']['coef_rel_change_pct']:.1f}%), with sign and statistical significance "
        f"preserved (Supplementary Table S10). This adjusted contrast is closer to applied epidemiological practice "
        "than unadjusted mean differences alone."
    )

    protocol_guide = (
        "Protocol application guide. Researchers adapting this blueprint to another survey should: "
        "(i) choose a reference cohort with complete contextual labels for internal validation; "
        "(ii) specify MCAR masking rates matching expected incompleteness and at least one MAR mechanism "
        "tied to observed covariates with reported propensity coefficients; "
        "(iii) pre-declare a primary estimand that is not defined by the same features used as predictors "
        "(prefer administrative or external labels for the primary task); "
        "(iv) evaluate both predictive metrics and at least one downstream target parameter "
        "(e.g., an adjusted exposure coefficient) under true versus imputed labels; "
        "(v) assess temporal and/or spatial holdout generalisability; and "
        "(vi) report calibration if probabilities will be reused as weights or multiple-imputation draws. "
        "Hyperparameters and comparator imputation methods can be swapped without changing the protocol skeleton."
    )

    miss_results = (
        f"Under the explicit mechanisms above, proposed accuracy was {m['mcar_acc']:.3f} (MCAR), "
        f"{m['mar_acc']:.3f} (MAR), and {m['spat_acc']:.3f} (spatial rates). The larger drop under MAR "
        "indicates sensitivity when missingness tracks FatER—the same feature that enters prediction—"
        "reinforcing reporting of mechanism parameters rather than MCAR-only claims."
    )

    # Replace / insert content
    replacements = [
        ("We formalise a reusable statistical evaluation protocol", pillars),
        ("The present study addresses this gap by proposing a reproducible statistical evaluation protocol", pillars),
        ("This manuscript extends our prior conference work", pillars),
        ("Primary evaluations masked contextual labels completely at random", missing_section),
        ("MCAR. On the held-out test set, each observation's contextual label is masked", missing_section),
        ("Calibration is evaluated because well-calibrated class probabilities", cal_rationale),
        ("Predicted probabilities were reasonably calibrated", cal_rationale),
        ("Predicted probabilities were well calibrated", cal_rationale),
        ("Jensen–Shannon divergence summarises agreement", js_rationale),
        ("Distributional agreement between observed and inferred", js_rationale),
        ("Threshold fragility of the three-class descriptive labels", s9_results),
        ("Covariate-adjusted downstream preservation", adj_results),
        ("Protocol application guide", protocol_guide),
        ("Under the explicit mechanisms above, proposed accuracy was", miss_results),
    ]

    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        for prefix, new in replacements:
            if p.text.startswith(prefix) or prefix in p.text[:80]:
                # Only replace if matching intended paragraph types
                if prefix in p.text:
                    replace_para(p, new)
                    break

    # Ensure pillars exist
    if not any("three pillars" in p.text for p in doc.paragraphs):
        insert_before(doc, "We propose a gradient-boosted", pillars)
        if not any("three pillars" in p.text for p in doc.paragraphs):
            insert_before(doc, "Study Design and Data Source", pillars)

    # Missingness heading + body
    if not any(p.text.strip() == "Missing Data Mechanisms and Simulation" for p in doc.paragraphs):
        insert_heading_before(doc, "MCAR. On the held-out test set", "Missing Data Mechanisms and Simulation")
        if not any(p.text.strip() == "Missing Data Mechanisms and Simulation" for p in doc.paragraphs):
            insert_heading_before(
                doc,
                "Primary evaluations masked contextual labels",
                "Missing Data Mechanisms and Simulation",
            )

    if not any("logit{P(R=1 | FatER)}" in p.text or "logit{{P(R=1 | FatER)}}" in p.text or "β₁ =" in p.text for p in doc.paragraphs):
        # missing_section may use special chars - check beta
        if not any("beta1" in p.text.lower() or "β1" in p.text or "β₁" in p.text for p in doc.paragraphs):
            insert_before(doc, "To address class imbalance", missing_section)

    if not any("inverse-probability weights" in p.text for p in doc.paragraphs):
        insert_before(doc, "Analyses using inferred labels preserved", cal_rationale)
        insert_before(doc, "Analyses using imputed labels preserved", cal_rationale)

    if not any("necessary but not sufficient" in p.text for p in doc.paragraphs):
        insert_before(doc, "Leave-one-year-out validation demonstrated", js_rationale)

    if not any("Threshold fragility of the three-class" in p.text for p in doc.paragraphs):
        insert_before(doc, "Supplementary analysis: Binary administrative", s9_results)
        if not any("Threshold fragility" in p.text for p in doc.paragraphs):
            insert_before(doc, "Robustness Across Missing Rates", s9_results)

    if not any("Covariate-adjusted downstream preservation" in p.text for p in doc.paragraphs):
        insert_before(doc, "Supplementary Analyses", adj_results)

    if not any("Protocol application guide" in p.text for p in doc.paragraphs):
        insert_before(doc, "### Strengths", protocol_guide)
        if not any("Protocol application guide" in p.text for p in doc.paragraphs):
            insert_before(doc, "Several limitations should be considered", protocol_guide)

    if not any("Under the explicit mechanisms above" in p.text for p in doc.paragraphs):
        insert_before(doc, "Threshold fragility of the three-class", miss_results)

    # Terminology: light touch on key repeated phrases (label recovery)
    for p in doc.paragraphs:
        if has_equation(p._element):
            continue
        t = p.text
        # Avoid changing "statistical inference" / "inferential target" / "downstream inferential"
        if "inferential" in t or "Inference Framework" in t or "statistical inference" in t:
            continue
        if "inferred labels" in t:
            replace_para(p, t.replace("inferred labels", "imputed labels"))
        elif "label inference" in t:
            replace_para(p, t.replace("label inference", "label imputation"))

    remove_duplicates(doc, "We formalise a reusable statistical evaluation protocol")
    remove_duplicates(doc, "MCAR. On the held-out test set")
    remove_duplicates(doc, "Threshold fragility of the three-class")
    remove_duplicates(doc, "Protocol application guide")
    remove_duplicates(doc, "Covariate-adjusted downstream preservation")
    remove_duplicates(doc, "Calibration is evaluated because")
    remove_duplicates(doc, "Jensen–Shannon divergence summarises")

    # Data/code availability nudge in conclusions area
    avail = (
        "Code and data availability. Analysis scripts that regenerate all main and supplementary tables "
        "(including Table S9 band sensitivity and Table S10 adjusted regressions) are provided at "
        "https://github.com/tyrone1979/nutri2-urban."
    )
    if not any("Code and data availability" in p.text for p in doc.paragraphs):
        insert_before(doc, "Future work should prioritise", avail)

    doc.save(str(MAIN))
    print(f"Updated {MAIN}")


def apply_supplementary(m):
    # Rebuild base supp then append S10 + improved Fig S1 note via paragraph edits
    import subprocess
    import sys

    r = subprocess.run([sys.executable, "-u", str(ROOT / "update_supplementary_docx.py")], cwd=ROOT)
    if r.returncode != 0:
        raise SystemExit("update_supplementary_docx failed")

    doc = Document(str(SUPP))
    # Improve Fig S1 caption
    for p in doc.paragraphs:
        if p.text.startswith("Figure S1. SHAP summary"):
            replace_para(
                p,
                "Figure S1. SHAP summary plot for Urban-category contributions. Each point is one observation; "
                "horizontal position is the SHAP value (impact on the Urban log-odds scale). Colour encodes the "
                "raw feature value from low (blue) to high (red). Points are vertically jittered within each "
                "feature row to reduce overplotting (not ordered by feature value).",
            )
        if "Transitional accuracy is omitted" in p.text:
            replace_para(
                p,
                "Transitional accuracy is omitted because that stratum is partially defined by FatER and yields "
                "near-perfect scores that are not informative for administrative label recovery. Urban accuracy "
                "remains the challenging stratum (~0.41–0.43). Band-threshold fragility of three-class accuracy "
                "is reported in Table S9 and discussed as a primary reason for privileging the binary task.",
            )

    # Append S10 if missing
    if not any("Supplementary Table S10" in p.text for p in doc.paragraphs):
        from docx.shared import Pt as Pt2

        p = doc.add_paragraph(
            "Supplementary Table S10. Covariate-adjusted Urban versus Rural coefficients "
            "(OLS; covariates: survey year and province indicators) under true versus imputed labels "
            "(30% MCAR; Transitional excluded)."
        )
        for run in p.runs:
            run.bold = True
        doc.add_paragraph(
            "Age and sex are unavailable in the CHNS diet-file extract used here; Year and Province are the "
            "available adjustment set. Relative change = |coef_imp − coef_true| / |coef_true| × 100."
        )
        adj = m["adj"]
        table = doc.add_table(rows=1 + len(adj), cols=8)
        hdr = [
            "Outcome", "True coef", "True SE", "Imputed coef", "Imputed SE",
            "Δcoef (%)", "Sign OK", "Sig OK",
        ]
        for j, h in enumerate(hdr):
            table.rows[0].cells[j].paragraphs[0].text = h
        for i, (_, r) in enumerate(adj.iterrows(), start=1):
            vals = [
                r["outcome"],
                f"{r['true_coef']:.4f}",
                f"{r['true_se']:.4f}",
                f"{r['imp_coef']:.4f}",
                f"{r['imp_se']:.4f}",
                f"{r['coef_rel_change_pct']:.1f}",
                "Yes" if r["sign_consistent"] else "No",
                "Yes" if r["sig_consistent"] else "No",
            ]
            for j, v in enumerate(vals):
                table.rows[i].cells[j].paragraphs[0].text = v

    doc.save(str(SUPP))
    print(f"Updated {SUPP}")


def write_response(m):
    text = f"""# Response to SiM Major/Minor Revisions

## Major 1 — Elevating methodological contribution
- Reframed Introduction around **three protocol pillars**: (1) simulated missingness internal validation; (2) separating predictive accuracy from downstream inferential preservation; (3) spatiotemporal generalisability.
- Added a **Protocol application guide** in the Discussion for reuse on other surveys.
- Contribution positioned as a reusable evaluation blueprint, with CHNS as the illustrative testbed (not as an XGBoost methods paper).

## Major 2 — Missingness mechanisms
- New subsection **Missing Data Mechanisms and Simulation**.
- **MCAR**: independent Bernoulli masking with π = 0.30 on the held-out test set.
- **MAR**: logit{{P(R=1|FatER)}} = β₀ + β₁ z(FatER) with β₁ = {m['mar']['beta1']:.1f}, β₀ = {m['mar']['beta0']:.3f}, realised rate {m['mar']['realized_rate']:.3f}.
- **Spatial**: observation-level rates 0.50 (Beijing/Shanghai/Chongqing) vs 0.20 otherwise (not province deletion).
- Parameters saved in `results/missingness_mechanism_params.csv`; scripts: `missingness_simulation.py`.

## Major 3 — Metric rationale
- Calibration: framed as enabling **IPW / multiple imputation** reuse of class probabilities, not only fit diagnostics.
- JS divergence: stated as **necessary but not sufficient** (margins ≠ joints); temporal/spatial/downstream checks retained.

## Major 4 — Downstream adjusted analysis
- New analysis: OLS of macronutrients on Urban vs Rural, adjusting for **Year + Province** (age/sex unavailable in diet extract).
- FatER Urban coefficient: true {m['fat']['true_coef']:.4f} (SE {m['fat']['true_se']:.4f}) vs imputed {m['fat']['imp_coef']:.4f} (SE {m['fat']['imp_se']:.4f}); relative change {m['fat']['coef_rel_change_pct']:.1f}% (Supplementary Table S10).
- Script: `downstream_adjusted_regression.py`.

## Major 5 — Table S9 fragility → primary binary task
- Results now highlight accuracy drop from {m['band_default']['accuracy']:.3f} ([0.23,0.30]) to {m['band_wide']['accuracy']:.3f} ([0.25,0.32]).
- Discussed as evidence that predictor-defined hard thresholds make multi-class accuracy unstable; binary T2 task remains primary (accuracy {m['binary']['accuracy_masked']:.3f}).

## Minor
- Fig S1 caption expanded (colour scale; vertical jitter).
- Terminology: label recovery phrased as **imputation/prediction**; “inference” reserved for parameter estimation / protocol pillars where appropriate.
- GitHub (`tyrone1979/nutri2-urban`) includes scripts regenerating S9/S10 and mechanism parameters.
"""
    RESPONSE.write_text(text, encoding="utf-8")
    print(f"Wrote {RESPONSE}")


def main():
    m = load_metrics()
    apply_main(m)
    apply_supplementary(m)
    write_response(m)


if __name__ == "__main__":
    main()
